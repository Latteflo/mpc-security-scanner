"""
Word Document (DOCX) Report Generation
Professional Microsoft Word reports
"""

import sys
from pathlib import Path
from typing import List
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from models import MCPServer, Vulnerability, Severity
from utils.logger import get_logger

logger = get_logger("word_reporter")


class WordReportGenerator:
    """Generates Microsoft Word (DOCX) reports"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        self.severity_colors = {
            Severity.CRITICAL: RGBColor(231, 76, 60),    # Red
            Severity.HIGH: RGBColor(230, 126, 34),       # Orange
            Severity.MEDIUM: RGBColor(241, 196, 15),     # Yellow
            Severity.LOW: RGBColor(52, 152, 219),        # Blue
            Severity.INFO: RGBColor(39, 174, 96)         # Green
        }
    
    def generate(self, server: MCPServer, vulnerabilities: List[Vulnerability],
                output_path: str) -> str:
        """Generate Word document report"""
        
        logger.info(f"Generating Word report to {output_path}")
        
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "MCP Security Assessment Report"
        doc.core_properties.subject = f"Security scan of {server.url}"
        doc.core_properties.author = "MCP Security Scanner"
        doc.core_properties.comments = f"Generated on {datetime.now().strftime('%Y-%m-%d')}"
        
        # Add content
        self._add_cover_page(doc, server)
        self._add_page_break(doc)
        
        self._add_executive_summary(doc, server, vulnerabilities)
        self._add_page_break(doc)
        
        self._add_findings(doc, vulnerabilities)
        self._add_page_break(doc)
        
        self._add_recommendations(doc)
        
        # Save document
        doc.save(str(output_file))
        
        logger.info(f"Word report saved to {output_file}")
        return str(output_file)
    
    def _add_cover_page(self, doc: Document, server: MCPServer):
        """Add cover page"""
        
        # Title
        title = doc.add_heading('', level=0)
        title_run = title.add_run('🔒 SECURITY ASSESSMENT REPORT')
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f'Model Context Protocol Server\n{server.url}')
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(127, 140, 141)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        for _ in range(5):
            doc.add_paragraph()
        
        # Info table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info = [
            ('Report Date:', datetime.now().strftime('%B %d, %Y')),
            ('Target Server:', server.url),
            ('Server Name:', server.name or 'Unknown'),
            ('Assessment Type:', 'Comprehensive Security Scan'),
            ('Status:', 'CONFIDENTIAL')
        ]
        
        for i, (label, value) in enumerate(info):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_executive_summary(self, doc: Document, server: MCPServer,
                               vulnerabilities: List[Vulnerability]):
        """Add executive summary"""
        
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        
        # Summary text
        summary = doc.add_paragraph()
        summary.add_run(
            f"This report presents the findings from a comprehensive security assessment "
            f"of the Model Context Protocol (MCP) server deployed at {server.url}. "
            f"The assessment identified {len(vulnerabilities)} security issues "
            f"that require attention to ensure the confidentiality, integrity, and "
            f"availability of the system."
        )
        summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        
        # Calculate risk score
        severity_counts = {'CRITICAL': 0, 'HIGH': 0, 'MEDIUM': 0, 'LOW': 0, 'INFO': 0}
        for vuln in vulnerabilities:
            severity_counts[vuln.severity.value] += 1
        
        risk_score = min(sum([
            severity_counts['CRITICAL'] * 10,
            severity_counts['HIGH'] * 5,
            severity_counts['MEDIUM'] * 2,
            severity_counts['LOW'] * 1
        ]), 100)
        
        # Risk score box
        risk_para = doc.add_paragraph()
        risk_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        risk_run = risk_para.add_run(f'\nOVERALL RISK SCORE: {risk_score}/100\n')
        risk_run.font.size = Pt(24)
        risk_run.font.bold = True
        
        if risk_score >= 70:
            risk_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
            risk_level = "HIGH RISK"
        elif risk_score >= 40:
            risk_run.font.color.rgb = RGBColor(241, 196, 15)  # Yellow
            risk_level = "MEDIUM RISK"
        else:
            risk_run.font.color.rgb = RGBColor(39, 174, 96)  # Green
            risk_level = "LOW RISK"
        
        level_run = risk_para.add_run(f'{risk_level}\n')
        level_run.font.size = Pt(16)
        level_run.font.bold = True
        
        doc.add_paragraph()
        
        # Vulnerability breakdown table
        doc.add_heading('Vulnerability Breakdown', level=2)
        
        table = doc.add_table(rows=6, cols=4)
        table.style = 'Medium Grid 1 Accent 1'
        
        # Header
        headers = ['SEVERITY', 'COUNT', 'RISK LEVEL', 'PRIORITY']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Data
        data = [
            ('Critical', severity_counts['CRITICAL'], '🔴 Immediate', 'P0 - Fix Now'),
            ('High', severity_counts['HIGH'], '🟠 Urgent', 'P1 - This Week'),
            ('Medium', severity_counts['MEDIUM'], '🟡 Important', 'P2 - This Month'),
            ('Low', severity_counts['LOW'], '🔵 Minor', 'P3 - Backlog'),
            ('Info', severity_counts['INFO'], '🟢 FYI', 'P4 - Optional'),
        ]
        
        for i, row_data in enumerate(data, 1):
            row = table.rows[i]
            for j, value in enumerate(row_data):
                row.cells[j].text = str(value)
    
    def _add_findings(self, doc: Document, vulnerabilities: List[Vulnerability]):
        """Add detailed findings"""
        
        doc.add_heading('DETAILED FINDINGS', level=1)
        
        # Sort by severity
        severity_order = {Severity.CRITICAL: 0, Severity.HIGH: 1,
                         Severity.MEDIUM: 2, Severity.LOW: 3, Severity.INFO: 4}
        sorted_vulns = sorted(vulnerabilities,
                             key=lambda v: severity_order[v.severity])
        
        for i, vuln in enumerate(sorted_vulns, 1):
            self._add_vulnerability(doc, i, vuln)
            doc.add_paragraph()  # Spacing
    
    def _add_vulnerability(self, doc: Document, num: int, vuln: Vulnerability):
        """Add single vulnerability"""
        
        # Title with number
        title = doc.add_heading(f'{num}. {vuln.title}', level=2)
        title_run = title.runs[0]
        title_run.font.color.rgb = self.severity_colors.get(vuln.severity, RGBColor(0, 0, 0))
        
        # Metadata table
        table = doc.add_table(rows=3, cols=4)
        table.style = 'Light List Accent 1'
        
        meta = [
            ('ID:', vuln.id, 'Severity:', vuln.severity.value),
            ('Category:', vuln.category, 'CWE:', vuln.cwe_id or 'N/A'),
            ('CVSS:', f'{vuln.cvss_score}/10.0' if vuln.cvss_score else 'N/A', '', '')
        ]
        
        for i, row_data in enumerate(meta):
            row = table.rows[i]
            for j in range(0, len(row_data), 2):
                if j < len(row_data):
                    row.cells[j].text = row_data[j]
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
                if j+1 < len(row_data):
                    row.cells[j+1].text = str(row_data[j+1])
                    # Color severity
                    if row_data[j] == 'Severity:':
                        row.cells[j+1].paragraphs[0].runs[0].font.color.rgb = \
                            self.severity_colors.get(vuln.severity)
        
        # Description
        doc.add_paragraph()
        desc_heading = doc.add_paragraph()
        desc_heading.add_run('Description:').bold = True
        doc.add_paragraph(vuln.description)
        
        # Evidence
        doc.add_paragraph()
        evidence_heading = doc.add_paragraph()
        evidence_heading.add_run('Evidence:').bold = True
        for evidence in vuln.evidence[:5]:
            doc.add_paragraph(f'• {evidence}', style='List Bullet')
        
        # Remediation
        doc.add_paragraph()
        remed_heading = doc.add_paragraph()
        remed_heading.add_run('Remediation:').bold = True
        
        # Split remediation by newlines and add as bullet points
        for line in vuln.remediation.split('\n'):
            if line.strip():
                if line.strip().startswith('-'):
                    doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line.strip())
    
    def _add_recommendations(self, doc: Document):
        """Add recommendations section"""
        
        doc.add_heading('RECOMMENDATIONS', level=1)
        
        sections = [
            ('Immediate Actions (Priority 0):', [
                'Address all CRITICAL vulnerabilities within 24 hours',
                'Implement temporary mitigations if permanent fixes require time',
                'Notify security team and stakeholders'
            ]),
            ('Short-term Actions (1-2 weeks):', [
                'Resolve all HIGH severity issues',
                'Begin addressing MEDIUM severity vulnerabilities',
                'Implement monitoring and alerting'
            ]),
            ('Long-term Actions (1-3 months):', [
                'Address remaining MEDIUM and LOW severity issues',
                'Implement security best practices',
                'Schedule regular security assessments',
                'Provide security training to development team'
            ])
        ]
        
        for heading, items in sections:
            para = doc.add_paragraph()
            para.add_run(heading).bold = True
            
            for item in items:
                doc.add_paragraph(f'• {item}', style='List Bullet')
            
            doc.add_paragraph()  # Spacing
        
        # Disclaimer
        doc.add_paragraph()
        doc.add_paragraph()
        
        disclaimer = doc.add_paragraph()
        disclaimer.add_run('Disclaimer: ').bold = True
        disclaimer.add_run(
            'This report is provided for informational purposes only. '
            'The findings represent potential security issues identified through '
            'automated and manual testing. Manual verification is recommended before '
            'taking remediation actions.'
        )
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Footer
        footer = doc.add_paragraph()
        footer.add_run(f'\nReport Generated by: MCP Security Scanner v0.2.0\n')
        footer.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        footer.runs[0].font.size = Pt(9)
        footer.runs[1].font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_page_break(self, doc: Document):
        """Add page break"""
        doc.add_page_break()


# Convenience function
def generate_word_report(server: MCPServer, vulnerabilities: List[Vulnerability],
                        output_path: str) -> str:
    """Generate Word report - convenience function"""
    generator = WordReportGenerator()
    return generator.generate(server, vulnerabilities, output_path)
